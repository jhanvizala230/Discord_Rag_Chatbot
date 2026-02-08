import os
import tempfile

import camelot
import fitz
import pandas as pd
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend_langchain.core.logger import setup_logger
from backend_langchain.core.config import CHUNK_SIZE, CHUNK_OVERLAP
from backend_langchain.vector_and_db.table_file_store import store_table_as_json, store_table_as_csv

logger = setup_logger(__name__)


class DocumentExtractor:
    """Unified document extractor for PDF and DOCX files"""

    def __init__(
        self,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        camelot_flavors=None,
        min_image_dim=300,
    ):
        """
        Initialize the document extractor.
        
        Args:
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks
            camelot_flavors: Tuple of flavors for Camelot table detection (default: stream)
            min_image_dim: Minimum image dimension to extract (in pixels)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.camelot_flavors = camelot_flavors or ("stream",)
        self.min_image_dim = min_image_dim
        self.splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def extract(self, path, file_type):
        """
        Extract content from document based on file type.
        
        Args:
            path: Path to the document file
            file_type: Type of file ('pdf' or 'docx')
            
        Returns:
            dict with text_chunks, table_chunks, image_chunks, raw_tables
        """
        if file_type.lower() == "pdf":
            return self._extract_pdf(path)
        elif file_type.lower() == "docx":
            return self._extract_docx(path)
        else:
            raise ValueError(f"Unsupported file_type: {file_type}")

    def _extract_pdf(self, path):
        """Extract text, tables, and images from PDF"""
        logger.info(f"DocumentExtractor: Extracting PDF from {path}")
        text_sections = []
        text_chunks = []
        table_chunks = []
        image_chunks = []
        raw_tables = []
        source_name = os.path.basename(path)

        # Extract text and images using PyMuPDF
        try:
            with fitz.open(path) as pdf_doc:
                for page_index, page in enumerate(pdf_doc, start=1):
                    page_text = page.get_text("text")
                    if page_text and page_text.strip():
                        text_sections.append(page_text)
                    
                    # image extraction

                    for image_info in page.get_images(full=True):
                        xref = image_info[0]
                        pix = fitz.Pixmap(pdf_doc, xref)
                        width, height = pix.width, pix.height
                        if width >= self.min_image_dim and height >= self.min_image_dim:
                            fd, tmp_img_path = tempfile.mkstemp(suffix=".png")
                            os.close(fd)
                            pix_to_save = pix
                            if pix.n >= 5:
                                pix_to_save = fitz.Pixmap(fitz.csRGB, pix)
                            pix_to_save.save(tmp_img_path)
                            if pix_to_save is not pix:
                                pix_to_save = None
                            caption = f"Image on page {page_index} (likely diagram or chart)"
                            image_chunks.append({
                                "text": caption,
                                "metadata": {
                                    "content_type": "image",
                                    "image_type": "diagram",
                                    "page": page_index,
                                    "image_path": tmp_img_path,
                                    "source": source_name
                                }
                            })
                        pix = None
        except Exception as exc:
            logger.exception("PDF extraction failed for %s: %s", path, exc)
            raise

        # Extract tables using Camelot
        tables = []
        for flavor in self.camelot_flavors:
            try:
                candidate_tables = camelot.read_pdf(
                    path,
                    pages="all",
                    flavor=flavor,
                    suppress_stdout=True,
                )
            except Exception as exc:
                logger.warning("Camelot (%s) failed for %s: %s", flavor, path, exc)
                continue
            if candidate_tables:
                tables = candidate_tables
                break

        # Process extracted tables
        for table in tables:
            page_num = int(table.page) if table.page and table.page.isdigit() else None
            df = table.df.applymap(lambda cell: str(cell).replace("\n", " ").strip())
            table_rows = df.values.tolist()
            
            store_table_as_csv(source_name, page_num or 0, table_rows)
            store_table_as_json(source_name, page_num or 0, table_rows)

            csv_str = df.to_csv(index=False, header=False)
            raw_tables.append({
                "csv": csv_str,
                "metadata": {
                    "content_type": "table_raw",
                    "page": page_num,
                    "source": source_name
                }
            })

            preview_rows = [
                " | ".join(str(cell).strip() for cell in row).strip()
                for row in table_rows[:3]
                if any(str(cell).strip() for cell in row)
            ]
            preview_text = preview_rows[0] if preview_rows else "Table data captured"
            summary = f"Table on page {page_num or '?'}: {preview_text[:200]}"
            table_chunks.append({
                "text": summary,
                "metadata": {
                    "content_type": "table_summary",
                    "page": page_num,
                    "source": source_name,
                    "raw_table_csv": csv_str
                }
            })

        # Split text into chunks
        if text_sections:
            text = "\n".join(text_sections)
            for chunk in self.splitter.split_text(text):
                text_chunks.append({
                    "text": chunk,
                    "metadata": {
                        "content_type": "text",
                        "page": None,
                        "source": source_name
                    }
                })

        return {
            "text_chunks": text_chunks,
            "table_chunks": table_chunks,
            "image_chunks": image_chunks,
            "raw_tables": raw_tables
        }

    def _extract_docx(self, path):
        """Extract text and tables from DOCX"""
        logger.info(f"DocumentExtractor: Extracting DOCX from {path}")
        doc = Document(path)
        text = "\n\n".join([p.text for p in doc.paragraphs])
        text_chunks = []
        table_chunks = []
        raw_tables = []
        source_name = os.path.basename(path)

        # Process text into chunks
        if text.strip():
            for chunk in self.splitter.split_text(text):
                text_chunks.append({
                    "text": chunk,
                    "metadata": {
                        "content_type": "text",
                        "page": None,
                        "source": source_name
                    }
                })

        # Process tables
        for t_idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            
            df = pd.DataFrame(data[1:], columns=data[0] if data else [])
            csv_str = df.to_csv(index=False)
            
            store_table_as_csv(source_name, t_idx, data)
            store_table_as_json(source_name, t_idx, data)

            raw_tables.append({
                "csv": csv_str,
                "metadata": {
                    "content_type": "table_raw",
                    "page": t_idx,
                    "source": source_name
                }
            })

            summary = f"Table {t_idx}: {csv_str[:200]}..."
            table_chunks.append({
                "text": summary,
                "metadata": {
                    "content_type": "table_summary",
                    "page": t_idx,
                    "source": source_name,
                    "raw_table_csv": csv_str
                }
            })

        return {
            "text_chunks": text_chunks,
            "table_chunks": table_chunks,
            "image_chunks": [],
            "raw_tables": raw_tables
        }


def extract_and_chunk(path, file_type):
    """
    Unified extraction and chunking for PDF/DOCX files.
    This function is called by the document upload API.
    
    Args:
        path: Path to the document file
        file_type: Type of file ('pdf' or 'docx')
        
    Returns:
        dict: {"text_chunks": [...], "table_chunks": [...], "image_chunks": [...], "raw_tables": [...]}
    """
    try:
        extractor = DocumentExtractor()
        extracted = extractor.extract(path, file_type)
        return extracted
    except Exception as e:
        logger.error(f"extract_and_chunk_failed | error={str(e)} | path={path} | file_type={file_type}")
        raise

